from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re

class DocumentReader:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)
        
    def read_paragraphs(self):
        """Lê todos os parágrafos e elementos do documento incluindo imagens"""
        elements = []
        element_index = 0
        
        # Primeiro, processa parágrafos normais
        for i, para in enumerate(self.document.paragraphs):
            # Verifica se o parágrafo contém imagem inline
            has_inline_image = False
            for run in para.runs:
                if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                    has_inline_image = True
                    print(f"  Imagem inline detectada no parágrafo {i}")
                    break
            
            # Detecção detalhada de listas
            is_list_item = False
            list_type = None
            list_char = None
            
            # Verifica se é um item de lista formatado pelo Word
            if para._element.xpath('.//w:numPr'):
                is_list_item = True
                
                # Tenta identificar o tipo baseado no texto
                text_start = para.text.strip()[:10] if para.text else ""
                
                # Detecta tipo de marcador
                if text_start:
                    # Lista com letras (a), b), A), B)
                    if re.match(r'^[a-zA-Z][\)\.]\s', text_start):
                        list_type = 'letter'
                        list_char = text_start[0]
                    # Lista numerada 1. 2. 1) 2)
                    elif re.match(r'^\d+[\)\.]\s', text_start):
                        list_type = 'number'
                    # Bullets (•, -, *, etc)
                    else:
                        list_type = 'bullet'
                        # Tenta identificar o caractere do bullet
                        if para._element.xpath('.//w:lvlText'):
                            list_char = 'bullet'
            
            # Verifica também manualmente se parece uma lista (caso não esteja formatada)
            elif para.text:
                text_start = para.text.strip()
                # Padrões manuais de lista
                if re.match(r'^[a-eA-E][\)\.]\s', text_start):
                    is_list_item = True
                    list_type = 'letter'
                    list_char = text_start[0].upper()
                elif re.match(r'^\d+[\)\.]\s', text_start):
                    is_list_item = True
                    list_type = 'number'
                elif text_start.startswith(('• ', '- ', '* ', '→ ', '▪ ')):
                    is_list_item = True
                    list_type = 'bullet'
                    list_char = text_start[0]
            
            # SEMPRE adiciona o parágrafo, mesmo se vazio
            elements.append({
                'index': element_index,
                'type': 'paragraph',
                'text': para.text,  # Pode ser vazio
                'original_para_index': i,
                'style': para.style.name if para.style else 'Normal',
                'runs': self._extract_runs(para),
                'has_image': has_inline_image,
                'is_image_paragraph': has_inline_image and not para.text.strip(),
                'is_list_item': is_list_item,  # Se é item de lista
                'list_type': list_type,  # Tipo: 'bullet', 'number', 'letter'
                'list_char': list_char,  # Caractere usado (A, B, 1, 2, •, etc)
                'markers': []
            })
            element_index += 1
        
        # Processa tabelas
        for i, table in enumerate(self.document.tables):
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(' | '.join(row_text))
            
            if table_text:
                elements.append({
                    'index': element_index,
                    'type': 'table',
                    'text': '\n'.join(table_text),
                    'original_element': table._element,
                    'style': 'Table',
                    'markers': []
                })
                element_index += 1
        
        print(f"\nTotal de elementos lidos: {len(elements)}")
        
        # Conta tipos de elementos
        types_count = {}
        for elem in elements:
            elem_type = elem['type']
            types_count[elem_type] = types_count.get(elem_type, 0) + 1
        
        for elem_type, count in types_count.items():
            print(f"  - {elem_type}: {count}")
        
        # Debug: mostra onde estão as imagens
        image_positions = [elem['index'] for elem in elements if elem['type'] == 'image']
        if image_positions:
            print(f"  Posições das imagens: {image_positions}")
        
        return elements
    
    def _extract_runs(self, paragraph):
        """Extrai informações de formatação dos runs"""
        runs = []
        for run in paragraph.runs:
            runs.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size.pt if run.font.size else None,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            })
        return runs
    
    def read_tables(self):
        """Lê todas as tabelas do documento"""
        tables = []
        for i, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append({
                'index': i,
                'data': table_data
            })
        return tables
    
    def get_document_info(self):
        """Retorna informações gerais do documento"""
        # Conta elementos
        paragraph_count = 0
        image_count = 0
        table_count = 0
        
        for element in self.document.element.body:
            if element.tag.endswith('p'):
                paragraph_count += 1
                # Verifica se tem imagem
                para = None
                for p in self.document.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para:
                    for run in para.runs:
                        if run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'):
                            image_count += 1
                            break
            elif element.tag.endswith('tbl'):
                table_count += 1
        
        return {
            'total_paragraphs': paragraph_count,
            'total_images': image_count,
            'total_tables': table_count,
            'total_sections': len(self.document.sections),
            'core_properties': {
                'author': self.document.core_properties.author,
                'created': str(self.document.core_properties.created) if self.document.core_properties.created else None,
                'modified': str(self.document.core_properties.modified) if self.document.core_properties.modified else None,
                'title': self.document.core_properties.title
            }
        }