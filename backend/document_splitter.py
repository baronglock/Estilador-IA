from docx import Document
from typing import List, Dict, Tuple
import re

class DocumentSplitter:
    def __init__(self):
        self.simulado_pattern = re.compile(r'Simulado\s+(\d+)', re.IGNORECASE)
        
    def split_simulados(self, document: Document) -> List[Dict]:
        """Divide o documento em simulados individuais"""
        print("\nDividindo documento em simulados...")
        simulados = []
        current_simulado = None
        current_content = []
        
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip()
            
            # Verifica se é início de novo simulado (mais flexível)
            match = re.search(r'Simulado\s+(\d+)', text, re.IGNORECASE)
            if match and ('estudo' in text.lower() or len(text) < 50):  # Evita falsos positivos
                # Salva simulado anterior se existir
                if current_simulado and current_content:
                    simulados.append({
                        'number': current_simulado,
                        'content': current_content,
                        'start_index': current_start,
                        'end_index': i - 1
                    })
                    print(f"  Simulado {current_simulado}: {len(current_content)} parágrafos")
                
                # Inicia novo simulado
                current_simulado = int(match.group(1))
                current_start = i
                current_content = [para]
            elif current_simulado:
                # Adiciona ao simulado atual
                current_content.append(para)
        
        # Adiciona último simulado
        if current_simulado and current_content:
            simulados.append({
                'number': current_simulado,
                'content': current_content,
                'start_index': current_start,
                'end_index': len(document.paragraphs) - 1
            })
            print(f"  Simulado {current_simulado}: {len(current_content)} parágrafos")
        
        # Remove duplicatas e ordena
        unique_simulados = {}
        for sim in simulados:
            if sim['number'] not in unique_simulados or len(sim['content']) > len(unique_simulados[sim['number']]['content']):
                unique_simulados[sim['number']] = sim
        
        simulados = list(unique_simulados.values())
        simulados.sort(key=lambda x: x['number'])
        
        print(f"  Total de simulados únicos: {len(simulados)}")
        return simulados
    
    def separate_questions_answers_by_style(self, simulado_content: List, document: Document) -> Tuple[Document, Document]:
        """Separa questões de gabaritos baseado nos estilos aplicados"""
        print(f"\n  Separando questões e gabaritos por estilo...")
        
        questions_doc = Document()
        answers_doc = Document()
        
        # Copia estilos para ambos documentos
        self._copy_styles(document, questions_doc)
        self._copy_styles(document, answers_doc)
        
        gabarito_count = 0
        question_count = 0
        
        for para in simulado_content:
            # Verifica se o parágrafo tem estilo de gabarito
            is_gabarito = False
            
            # Verifica pelo nome do estilo
            if para.style and para.style.name:
                style_name_lower = para.style.name.lower()
                if 'gabarito' in style_name_lower or 'answer' in style_name_lower:
                    is_gabarito = True
            
            # Também verifica pelo conteúdo (backup)
            if not is_gabarito:
                text_lower = para.text.lower()
                # Padrões típicos de gabarito
                gabarito_patterns = [
                    r'^[a-h]\d+\s*[–-]',  # D4 –, H1 -, etc
                    r'^resposta:',
                    r'^gabarito:',
                    r'^alternativa correta:',
                    r'^é [a-h]\d+ porque',  # É D4 porque...
                    r'^o professor deve',
                    r'^esta questão avalia',
                    r'^habilidade:'
                ]
                
                for pattern in gabarito_patterns:
                    if re.search(pattern, text_lower):
                        is_gabarito = True
                        break
            
            # Adiciona ao documento apropriado
            if is_gabarito:
                self._copy_paragraph_to_document(answers_doc, para)
                gabarito_count += 1
            else:
                self._copy_paragraph_to_document(questions_doc, para)
                question_count += 1
        
        print(f"    - Parágrafos de questões: {question_count}")
        print(f"    - Parágrafos de gabarito: {gabarito_count}")
        
        return questions_doc, answers_doc
    
    def create_split_documents(self, simulados: List[Dict], document: Document) -> Dict[str, Document]:
        """Cria documentos separados para cada simulado"""
        print("\nCriando documentos separados...")
        documents = {}
        
        for simulado in simulados:
            sim_num = simulado['number']
            print(f"\n  Processando Simulado {sim_num}...")
            
            # Separa questões e gabaritos baseado em estilos
            questions_doc, answers_doc = self.separate_questions_answers_by_style(
                simulado['content'], 
                document
            )
            
            # Salva os documentos
            documents[f'simulado_{sim_num}_questoes'] = questions_doc
            documents[f'simulado_{sim_num}_gabarito'] = answers_doc
        
        return documents
    
    def create_complete_document(self, document: Document) -> Document:
        """Cria documento completo com todos os simulados"""
        print("\nCriando documento completo...")
        
        # Para o documento completo, cria uma cópia limpa
        complete_doc = Document()
        
        # Copia todos os estilos
        self._copy_styles(document, complete_doc)
        
        # Copia todos os parágrafos
        para_count = 0
        for para in document.paragraphs:
            if para.text.strip():  # Ignora parágrafos vazios
                self._copy_paragraph_to_document(complete_doc, para)
                para_count += 1
        
        print(f"  Documento completo criado com {para_count} parágrafos")
        return complete_doc
    
    def _copy_styles(self, source_doc: Document, target_doc: Document):
        """Copia estilos de um documento para outro"""
        for style in source_doc.styles:
            try:
                if hasattr(style, 'name') and not style.builtin:
                    # Tenta criar o estilo customizado
                    target_style = target_doc.styles.add_style(
                        style.name, 
                        style.type
                    )
                    
                    # Copia propriedades básicas
                    target_style.hidden = style.hidden
                    target_style.quick_style = style.quick_style
                    if hasattr(style, 'priority'):
                        target_style.priority = style.priority
                    
                    # Copia formatação de fonte
                    if style.font.color and style.font.color.rgb:
                        target_style.font.color.rgb = style.font.color.rgb
                    if style.font.size:
                        target_style.font.size = style.font.size
                    target_style.font.bold = style.font.bold
                    target_style.font.italic = style.font.italic
            except:
                # Ignora erros (estilos built-in ou já existentes)
                pass
    
    def _copy_paragraph_to_document(self, target_doc: Document, source_para):
        """Copia um parágrafo mantendo estilo e formatação"""
        # Cria novo parágrafo
        new_para = target_doc.add_paragraph()
        
        # Copia o estilo
        try:
            if source_para.style:
                new_para.style = source_para.style.name
        except:
            pass
        
        # Copia runs com formatação
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            
            # Copia formatação
            if run.bold is not None:
                new_run.bold = run.bold
            if run.italic is not None:
                new_run.italic = run.italic
            if run.underline is not None:
                new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
        
        # Copia formatação de parágrafo
        if source_para.paragraph_format.left_indent:
            new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        if source_para.paragraph_format.space_before:
            new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        if source_para.paragraph_format.space_after:
            new_para.paragraph_format.space_after = source_para.paragraph_format.space_after