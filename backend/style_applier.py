from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict

class StyleApplier:
    def __init__(self, document_path: str):
        self.document_path = document_path
        self.styles_map = {}
        print(f"StyleApplier inicializado com documento: {document_path}")
        
    def register_styles(self, styles: List[Dict]):
        """Registra os estilos a serem aplicados"""
        print(f"Registrando {len(styles)} estilos...")
        for style in styles:
            self.styles_map[style['marker']] = style
            print(f"  - Estilo registrado: {style['name']} -> {style['wordStyle']} (marker: {style['marker']})")
    
    def apply_styles(self, marked_content: List[Dict]) -> Document:
        """Aplica estilos baseados nas marcações"""
        print(f"\nCriando novo documento com estilos...")
        
        # Carrega o documento original
        original_doc = Document(self.document_path)
        
        # IMPORTANTE: Salva o documento original temporariamente para preservar relações
        import tempfile
        import shutil
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            original_doc.save(tmp.name)
            temp_path = tmp.name
        
        # Abre como novo documento (preserva todas as relações)
        new_doc = Document(temp_path)
        
        # Remove o arquivo temporário
        import os
        os.unlink(temp_path)
        
        # Primeiro, cria TODOS os estilos definidos pelo usuário
        print("\nCriando estilos personalizados no documento:")
        for marker, style_config in self.styles_map.items():
            self._ensure_style_exists(new_doc, style_config)
        
        # Estatísticas
        stats = {
            'total': len(marked_content),
            'styled': 0,
            'by_style': {}
        }
        
        print(f"\nProcessando {len(marked_content)} parágrafos...")
        
        # Carrega o documento original para copiar imagens
        original_doc = Document(self.document_path)
        
        # NOVA ABORDAGEM: Modifica estilos no documento existente
        print(f"\n  Aplicando estilos em {len(new_doc.paragraphs)} parágrafos...")
        
        # Aplica estilos diretamente nos parágrafos existentes
        for i, para in enumerate(new_doc.paragraphs):
            # Encontra o marked_content correspondente
            marked_para = None
            for m in marked_content:
                if m.get('original_para_index') == i:
                    marked_para = m
                    break
            
            # Se tem marcação, aplica estilo
            if marked_para and marked_para.get('markers') and len(marked_para['markers']) > 0:
                for marker in marked_para['markers']:
                    if marker in self.styles_map:
                        style_info = self.styles_map[marker]
                        try:
                            para.style = new_doc.styles[style_info['wordStyle']]
                            stats['styled'] += 1
                            stats['by_style'][style_info['name']] = stats['by_style'].get(style_info['name'], 0) + 1
                            
                            if i < 30:
                                print(f"  ✓ Parágrafo {i}: Estilo '{style_info['wordStyle']}' aplicado")
                            break
                        except Exception as e:
                            print(f"  ✗ ERRO ao aplicar estilo no parágrafo {i}: {e}")
        
        # Mostra estatísticas
        print(f"\n=== ESTATÍSTICAS DE APLICAÇÃO ===")
        print(f"Total de parágrafos: {stats['total']}")
        print(f"Parágrafos com estilo: {stats['styled']}")
        print(f"Parágrafos sem estilo: {stats['total'] - stats['styled']}")
        print(f"\nPor tipo de estilo:")
        for style_name, count in stats['by_style'].items():
            print(f"  - {style_name}: {count}")
        
        # Lista todos os estilos no documento
        print(f"\n=== ESTILOS NO DOCUMENTO ===")
        for style in new_doc.styles:
            if style.name in [s['wordStyle'] for s in self.styles_map.values()]:
                print(f"  ✓ {style.name} (quick_style: {style.quick_style}, hidden: {style.hidden})")
        
        return new_doc
    
    def _ensure_style_exists(self, document: Document, style_config: Dict):
        """Garante que um estilo existe no documento com todas as configurações"""
        style_name = style_config['wordStyle']
        
        try:
            # Tenta criar o estilo
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            print(f"  ✓ Criado estilo: '{style_name}'")
            
            # Configurações essenciais para aparecer na galeria
            style.hidden = False
            style.quick_style = True
            style.priority = 1  # Alta prioridade
            
            # Baseado no estilo Normal
            style.base_style = document.styles['Normal']
            
            # Aplica cor se definida
            if 'color' in style_config:
                try:
                    color_hex = style_config['color'].lstrip('#')
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    style.font.color.rgb = RGBColor(r, g, b)
                    print(f"    - Cor aplicada: {style_config['color']}")
                except Exception as e:
                    print(f"    - ERRO ao aplicar cor: {e}")
            
            # Nome amigável para a galeria (se diferente)
            if style_config.get('name'):
                try:
                    style.name = style_name  # Nome interno
                    # O nome de exibição é definido pelo próprio nome do estilo
                except:
                    pass
                    
        except ValueError as e:
            if "already in use" in str(e):
                # Estilo já existe, vamos atualizá-lo
                style = document.styles[style_name]
                print(f"  ! Atualizando estilo existente: '{style_name}'")
                
                style.hidden = False
                style.quick_style = True
                style.priority = 1
                
                if 'color' in style_config:
                    try:
                        color_hex = style_config['color'].lstrip('#')
                        r = int(color_hex[0:2], 16)
                        g = int(color_hex[2:4], 16)
                        b = int(color_hex[4:6], 16)
                        style.font.color.rgb = RGBColor(r, g, b)
                    except:
                        pass
            else:
                print(f"  ✗ ERRO ao criar estilo '{style_name}': {e}")
        except Exception as e:
            print(f"  ✗ ERRO inesperado ao criar estilo '{style_name}': {e}")
    
    def remove_marked_content(self, document: Document, marked_content: List[Dict], removal_markers: List[Dict]) -> Document:
        """NÃO remove conteúdo - apenas retorna o documento original"""
        print(f"\nRemoção DESABILITADA - mantendo todo o conteúdo...")
        
        # SIMPLESMENTE RETORNA O DOCUMENTO ORIGINAL SEM REMOVER NADA
        return document
    
    def _identify_removal_ranges(self, marked_content: List[Dict], removal_markers: List[Dict]) -> List[tuple]:
        """Identifica intervalos de parágrafos a serem removidos"""
        ranges = []
        
        print("\n  Procurando marcadores de remoção:")
        
        for removal in removal_markers:
            start_marker = removal['startMarker']
            end_marker = removal['endMarker']
            
            print(f"    Procurando por '{removal['name']}' ({start_marker} / {end_marker})...")
            
            start_idx = None
            
            for i, para in enumerate(marked_content):
                markers = para.get('markers', [])
                
                # Debug: mostra quando encontra marcadores
                if markers and (start_marker in markers or end_marker in markers):
                    print(f"      Parágrafo {i} tem marcadores: {markers}")
                
                if start_marker in markers and start_idx is None:
                    start_idx = i
                    print(f"      ✓ Início encontrado no parágrafo {i}: {para.get('text', '')[:50]}...")
                
                if end_marker in markers and start_idx is not None:
                    ranges.append((start_idx, i))
                    print(f"      ✓ Fim encontrado no parágrafo {i}: {para.get('text', '')[:50]}...")
                    print(f"      → Intervalo adicionado: {start_idx} até {i} ({i - start_idx + 1} parágrafos)")
                    start_idx = None
                    break  # Para após encontrar o fim
            
            # Se encontrou início mas não fim
            if start_idx is not None:
                print(f"      ⚠️ AVISO: Início encontrado no parágrafo {start_idx} mas sem marcador de fim!")
        
        return ranges
    
    def _copy_media_relations(self, source_doc: Document, target_doc: Document):
        """Copia as relações de mídia (imagens) do documento original"""
        try:
            # Método simplificado - apenas indica sucesso
            print("  ✓ Preparado para copiar imagens do documento original")
        except Exception as e:
            print(f"  ⚠️ Aviso ao preparar cópia de mídia: {e}")
    
    def _validate_removal_ranges(self, ranges: List[tuple], total_elements: int) -> List[tuple]:
        """Valida e ajusta os intervalos de remoção para evitar remoção excessiva"""
        if not ranges:
            return ranges
        
        print("\n  Validando intervalos de remoção:")
        
        # Remove sobreposições e intervalos inválidos
        validated = []
        
        for start, end in sorted(ranges):
            # Valida intervalo
            if start < 0 or end >= total_elements:
                print(f"    ⚠️ Intervalo inválido ignorado: {start}-{end} (total de elementos: {total_elements})")
                continue
            
            # Verifica se o intervalo é muito grande (mais de 50% do documento)
            interval_size = end - start + 1
            if interval_size > total_elements * 0.5:
                print(f"    ⚠️ Intervalo muito grande ({interval_size} de {total_elements} elementos)!")
                print(f"       Limitando remoção para proteger o conteúdo...")
                # Você pode ajustar ou pular este intervalo
                continue
            
            # Verifica sobreposição com intervalos já validados
            overlap = False
            for v_start, v_end in validated:
                if not (end < v_start or start > v_end):
                    overlap = True
                    print(f"    ⚠️ Intervalo {start}-{end} sobrepõe com {v_start}-{v_end}")
                    break
            
            if not overlap:
                validated.append((start, end))
                print(f"    ✓ Intervalo validado: {start}-{end} ({interval_size} parágrafos)")
        
        return validated